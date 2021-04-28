from model.resource import *
from utils.Utils import *
import csv
import os
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
class WebpageInfo:
    resources = set()
    downloadPath = Utils().getDownloadPath()
    id = None 
    def clearResources(self):
        self.resources.clear()

    def __init__(self):
        resources=set()

    def setId(self,id):
        self.id=id

    def setResource(self,res : Resource):
        self.resources.add(res)

    def extendResources(self,res):
        self.resources.update(res)
    
    def getResources(self):
        return self.resources

    def getID(self):
        return self.id

    def printResources(self):
        for res in self.resources:
            r =Resource()
            r= res
            r.printAll()
    #metodo che scrive in un file .csv le risorse trovate
    def toCSV(self,filename,startTime,docFileName,url):
        print("Writing: "+  str(filename)+'.csv')
        numberOfRows = 1
        css=0
        img=0
        video=0
        others=0
        a=0
        downloaded=0
        with open(self.downloadPath+os.path.sep+ str(filename)+'.csv' ,'w', newline='',encoding="utf-8") as out_f:
            writer = csv.writer(out_f, delimiter=';')
            writer.writerow(['ID','TAG NAME','URL','NOME FILE','NOME ATTUALE','TESTO ALT','HREF','TESTO NEL TAG','FORMATO','STATUS'])
            
            for res in self.resources:
                r = Resource()
                r = res
                alt =r.getAlt()
                status =r.getStatus()
                tagName=r.getTagName()
                if alt=='fromCss':
                    css+=1
                if status == str(200):
                    downloaded+=1
                if tagName =='a':
                    a+=1
                elif tagName =='img':
                    img+=1
                elif tagName =='video':
                    video+=1
                else:
                    others+=1
                writer.writerow([numberOfRows,tagName,r.getUrl(),r.getFileName(),r.getNewFilename(),alt,r.getHref(),r.getText(),r.getFormat(),status])
                numberOfRows+=1
        out_f.close()
        self.writeStatistics(filename,numberOfRows,downloaded,css,a,img,video,others,startTime,docFileName,url)
    
    def writeStatistics(self,filename,numberOfRows,downloaded,css,a,img,video,others,startTime,docFileName,url):
        cssRatio=css/numberOfRows*100
        htmlRes=numberOfRows- css
        htmlRatio = htmlRes/numberOfRows*100
        aRatio=a/numberOfRows*100
        imgRatio= img/numberOfRows*100
        videoRatio= video/numberOfRows*100
        othersRatio= others/numberOfRows*100
        duration = time.time() - startTime

        doc  = docx.Document(self.downloadPath+os.path.sep+"Report "+docFileName+'.docx')
        style = doc.styles['Normal']
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        heading = doc.add_paragraph()
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        obj= doc.add_paragraph()
        obj_run= obj.add_run("Statistiche web scraping: {} ".format(url))
        obj_run.bold=True
        obj1= doc.add_paragraph()
        obj1.add_run("Durata: {} secondi\n" .format(duration))
        obj1.add_run("Risorse ricercate: {}\n" .format(numberOfRows))
        obj1.add_run("Risorse con status 200: {}\n" .format(downloaded))
        obj1.add_run("Risorse trovate dall'analizzatore css:{} -> {:.2f} %\n" .format(str(css),cssRatio))
        obj1.add_run("Risorse trovate dall'analizzatore html:{} -> {:.2f} %\n" .format(str(htmlRes),htmlRatio))
        obj1.add_run("Risorse con tag a: {} -> {:.2f} %\n" .format(str(a),aRatio))
        obj1.add_run("Risorse con tag img: {} -> {:.2f} %\n" .format(str(img),imgRatio))
        obj1.add_run("Risorse con tag video: {} -> {:.2f} %\n" .format(str(video),videoRatio))
        obj1.add_run("Risorse con altri tag: {} -> {:.2f} %\n" .format(str(others),othersRatio))
        obj1.alignment = 0
        doc.save(self.downloadPath+os.path.sep+"Report "+docFileName+'.docx')
        
    def appendToDataset(self,netloc):
        if os.path.isfile('BigFile.csv'):
            with open(self.downloadPath+os.path.sep+ str("BigFile")+'.csv' ,'a', newline='',encoding="utf-8") as out_f:
                writer = csv.writer(out_f)
                for resource in self.getResources():
                    alt=resource.getAlt()
                    text = resource.getText()
                    filename = resource.getFileName()
                    if filename!=None and filename.strip()!='':
                        writer.writerow([15,netloc,filename])
                    if alt!=None and filename.strip()!='':
                        writer.writerow([15,netloc,alt])
                    if text!=None and text.strip()!='':
                        writer.writerow([15,netloc,text])
            out_f.close()
            return
        else:
            try:
                with open(self.downloadPath+os.path.sep+ str("BigFile")+'.csv' ,'w', newline='',encoding="utf-8") as out_f:
                    writer = csv.writer(out_f)
                    writer.writerow(['VALUE','HOSTNAME','TEXT'])
                    for resource in self.getResources():
                        alt=resource.getAlt()
                        text =resource.getText()
                        filename =resource.getFileName()
                        if filename!=None and filename.strip()!='':
                            writer.writerow([15,netloc,filename])
                        if alt!=None and alt.strip()!='':
                            writer.writerow([15,netloc,alt])
                        if text!=None and text.strip()!='':
                            writer.writerow([15,netloc,text])
                out_f.close()
            except:
                print("close current file")
                pass
            return
            
           

