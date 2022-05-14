import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from utils.Utils import *

class Statistics:
    downloadPath ="/Reports"
    url=None
    numberOfpages=None
    res=None
    downloaded=None
    cssRes =None
    cssRatio =None
    htmlRes = None
    htmlRatio = None
    aRes=None
    aRatio = None
    imgRes = None
    imgRatio = None
    videoRes = None
    videoRatio = None
    othersRes =None
    othersRatio=None
    duration=None



    def writeToDoc(self,docFileName,detailed):
            PATH = "Reports/Report "+docFileName+'.docx'
            doc =None
            if os.path.exists(PATH):
                doc = docx.Document(PATH)
            else:
                doc = docx.Document()
            print(PATH)
            style = doc.styles['Normal']
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            heading = doc.add_paragraph()
            heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            obj= doc.add_paragraph()
            obj_run= obj.add_run("Statistiche web scraping: {} ".format(self.getUrl()))
            obj_run.bold=True
            obj1= doc.add_paragraph()
            if detailed:
                obj1.add_run("Durata: {} secondi\n" .format(self.getDuration()))
                obj1.add_run("Risorse ricercate: {}\n" .format(self.getRes()))
                obj1.add_run("Risorse scaricate: {}\n" .format(self.getDownloaded()))
                obj1.add_run("Risorse trovate dall'analizzatore css:{} -> {:.2f} %\n" .format(str(self.getCssRes()),self.getCssRatio()))
                obj1.add_run("Risorse trovate dall'analizzatore html:{} -> {:.2f} %\n" .format(str(self.getHtmlRes()),self.getHtmlRatio()))
                obj1.add_run("Risorse con tag a: {} -> {:.2f} %\n" .format(str(self.getAres()),self.getAratio()))
                obj1.add_run("Risorse con tag img: {} -> {:.2f} %\n" .format(str(self.getImgRes()),self.getImgRatio()))
                obj1.add_run("Risorse con tag video: {} -> {:.2f} %\n" .format(str(self.getVideoRes()),self.getVideoratio()))
                obj1.add_run("Risorse con altri tag: {} -> {:.2f} %\n" .format(str(self.getOtherRes()),self.getOtherRatio()))
            else:
                obj1.add_run("Numero pagine esaminate: {} \n" .format(self.getNumberOfPages()))
                obj1.add_run("Durata totale: {} secondi\n" .format(self.getDuration()))
                obj1.add_run("Risorse totali ricercate: {}\n" .format(self.getRes()))
                obj1.add_run("Risorse totali scaricate: {}\n" .format(self.getDownloaded()))
                obj1.add_run("Risorse totali trovate dall'analizzatore css:{} -> {:.2f} %\n" .format(str(self.getCssRes()),self.getCssRatio()))
                obj1.add_run("Risorse totali trovate dall'analizzatore html:{} -> {:.2f} %\n" .format(str(self.getHtmlRes()),self.getHtmlRatio()))
                obj1.add_run("Risorse totali con tag a: {} -> {:.2f} %\n" .format(str(self.getAres()),self.getAratio()))
                obj1.add_run("Risorse totali con tag img: {}-> {:.2f}%\n" .format(str(self.getImgRes()),self.getImgRatio()))
                obj1.add_run("Risorse totali con tag video: {} -> {:.2f}%\n" .format(str(self.getVideoRes()),self.getVideoratio()))
                obj1.add_run("Risorse totali con altri tag: {} -> {:.2f}%\n" .format(str(self.getOtherRes()),self.getOtherRatio()))
            obj1.alignment = 0
            doc.save(PATH)






    #setters
    def setDownloaded(self,down):
        self.downloaded=down
    def setUrl(self,url):
        self.url=url
    def setNumberOfPages(self,numberOfpages):
        self.numberOfpages= numberOfpages
    def setRes(self,res):
        self.res=res
    def setCssRes(self,cssRes):
        self.cssRes=cssRes
    def setCssRatio(self,cssRatio):
        self.cssRatio=cssRatio
    def setHtmlRes(self,htmlRes):
        self.htmlRes=htmlRes
    def setHtmlRatio(self,htmlRatio):
        self.htmlRatio=htmlRatio
    def setAres(self,aRes):
        self.aRes=aRes
    def setAratio(self,aRatio):
        self.aRatio=aRatio
    def setImgRes(self,imgRes):
        self.imgRes=imgRes
    def setImgRatio(self,imgRatio):
        self.imgRatio=imgRatio
    def setVideoRes(self,videoRes):
        self.videoRes=videoRes
    def setVideoratio(self,videoRatio):
        self.videoRatio=videoRatio
    def setOtherRes(self,othersRes):
        self.othersRes=othersRes
    def setOtherRatio(self,othersRatio):
        self.othersRatio=othersRatio
    def setDuration(self,duration):
        self.duration=duration
    #getters
    def getUrl(self):
        return self.url
    def getRes(self):
        return self.res
    def getNumberOfPages(self):
        return self.numberOfpages
    def getDownloaded(self):
        return self.downloaded
    def getCssRes(self) :
        return self.cssRes
    def getCssRatio(self) :
        return self.cssRatio
    def getHtmlRes(self) :
        return self.htmlRes
    def getHtmlRatio(self) :
        return self.htmlRatio
    def getAres(self) :
        return self.aRes
    def getAratio(self) :
        return self.aRatio
    def getImgRes(self) :
        return self.imgRes
    def getImgRatio(self) :
        return self.imgRatio
    def getVideoRes(self) :
        return self.videoRes
    def getVideoratio(self) :
        return self.videoRatio
    def getOtherRes(self) :
        return self.othersRes
    def getOtherRatio(self) :
        return self.othersRatio
    def getDuration(self) :
        return self.duration