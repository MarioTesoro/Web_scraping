import docx                                     #libreria necessaria per gestire i file word
from docx.shared import RGBColor, Pt, Inches    #necessario per impostare colore del testo e grandezza
from docx.enum.text import WD_ALIGN_PARAGRAPH   #necessario per allineamento del testo
from Model.WhoisRecord import Whois             #necessario per gestire oggetti di tipo WHOIS
from Model.Malware import Malware               #necessario per gestire oggetti di tipo Malware
import collections
from ReportManagement.MalwareGraph import malware_graph
import os

#costante per elenco puntato
list_type = 'List Bullet 2'

def find_top_malware(dict_malware_info,list_malware_info):
    number = 0
    list_of_keys = dict_malware_info.keys()

    for key in list_of_keys:
       dict_malware_info[key] = list_malware_info.count(key)

    sorted_dict = collections.OrderedDict(dict_malware_info)
    sorted_list = sorted_dict.keys()
    return sorted_list, sorted_dict

#questo metodo si occupa di scrivere tutte le informazioni utili per il report
#riguardo un'URL, all'interno di un file word
def write_report(url, whois_info, malware_info, malware_number,category_type, category_description):

    #creazione oggetti Whois e Malware
    whois_obj = Whois()
    malware_obj = Malware()
    
    #suddivisione dell'URL ricevuto, da utilizzare nel titolo 
    #e nel nome del file
    url_splitted = url.split('/')
    #impostazione generali del documento
    doc  = docx.Document()
    style = doc.styles['Normal']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    #scrittura del titolo
    heading = doc.add_paragraph()
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_heading = heading.add_run(f'Report del sito {url_splitted[2]}')
    run_heading.font.size = Pt(24)
    run_heading.font.color.rgb = RGBColor.from_string('000000')
    run_heading.bold = True

    #spaziatura 
    doc.add_paragraph()

    #acquisizione record WHOIS per scrittua d'informazioni
    whois_obj = whois_info

    #scrittua del dominio e categoria, dominio preso dal record WHOIS e categoria tramite X-Force
    parag_domain = doc.add_paragraph()
    run_domain = parag_domain.add_run('Dominio: ')
    run_domain.bold = True
    run_site_domain = parag_domain.add_run(str(whois_obj.get_domain()))
    parag_category = doc.add_paragraph()
    run_category = parag_category.add_run('Categoria: ')
    run_category.bold = True
    #se la categoria esiste la scrive, altrimenti no
    if category_type != True:
        parag_category.add_run(str(category_type)+': '+str(category_description))

    #spaziatura
    doc.add_paragraph()

    #scrittura di tutti i campi del record WHOIS
    parag_whois = doc.add_paragraph()
    run_whois = parag_whois.add_run('RECORD WHOIS:')
    run_whois.bold = True


    parag_creation = doc.add_paragraph()
    run_creation = parag_creation.add_run('Data di Creazione: ')
    run_creation.bold = True
    parag_creation.style = list_type
    parag_creation.add_run(str(whois_obj.get_creation_date()))

    parag_update = doc.add_paragraph()
    run_update = parag_update.add_run('Data di Aggiornamento: ')
    run_update.bold = True
    parag_update.style = list_type
    parag_update.add_run(str(whois_obj.get_updated_date()))

    parag_expiration = doc.add_paragraph()
    run_expiration = parag_expiration.add_run('Data di Scadenza: ')
    run_expiration.bold = True
    parag_expiration.style = list_type
    parag_expiration.add_run(str(whois_obj.get_expiration_date()))

    parag_name = doc.add_paragraph()
    run_name = parag_name.add_run('Nome del registrante: ')
    run_name.bold = True
    parag_name.style = list_type
    parag_name.add_run(str(whois_obj.get_name()))

    parag_org = doc.add_paragraph()
    run_org = parag_org.add_run('Organizzazione del registrante: ')
    run_org.bold = True
    parag_org.style = list_type
    parag_org.add_run(str(whois_obj.get_org()))

    parag_location = doc.add_paragraph()
    run_location = parag_location.add_run('Città e/o regione del registrante: ')
    run_location.bold = True
    parag_location.style = list_type
    parag_location.add_run(str(whois_obj.get_location()))

    parag_registrar = doc.add_paragraph()
    run_registrar = parag_registrar.add_run('Nome del registrar: ')
    run_registrar.bold = True
    parag_registrar.style = list_type
    parag_registrar.add_run(str(whois_obj.get_registrar()))

    parag_email = doc.add_paragraph()
    run_email = parag_email.add_run('Email associate: ')
    run_email.bold = True
    parag_email.style = list_type
    parag_email.add_run(str(whois_obj.get_email()))

    #spaziatura
    doc.add_paragraph()

    parag_malware = doc.add_paragraph()
    run_malware = parag_malware.add_run('MALWARE')
    run_malware.bold = True

    #se sono stati trovati malware per quell'URL
    if malware_info != True:
        #scrive quanti ne sono stati trovati
        parag_num_malware = doc.add_paragraph()
        run_num = parag_num_malware.add_run('Numero malware trovati:')
        run_num.bold = True
        parag_num_malware.add_run(str(malware_number))
        list_type_malware = []
        dict_type_malware = {}
        malware_family = []
        dict_malware_family = {}
        parag_title_type = doc.add_paragraph()
        run_type = parag_title_type.add_run('TOP tipologie di Malware trovati: ')
        run_type.bold = True
        #scrive una TOP di tipologie e famiglie presenti nel sito, con relativo grafico
        for malware in malware_info:
            
            malware_obj = malware
            
            type_malware = str(malware_obj.get_malware_type())
            list_type_malware.append(type_malware)
            if type_malware not in dict_type_malware.keys():
                dict_type_malware[type_malware] = 0

            malware_family.extend(malware_obj.get_family())
            for family in malware_family:
                if family not in dict_malware_family.keys():
                    dict_malware_family[family] = 0

        list_number = 1
        sorted_malware,sorted_dict = find_top_malware(dict_type_malware,list_type_malware)
        file_name = malware_graph(sorted_dict,url_splitted[2],'type')

        for malware in sorted_malware:
            parag_type = doc.add_paragraph()
            run_list = parag_type.add_run('     '+str(list_number)+'. ')
            run_list.bold = True
            run_type = parag_type.add_run('Tipologia '+malware+' trovata '+str(sorted_dict[malware])+f' volte/a su {malware_number} totali')
            list_number += 1
        
        parag_image = doc.add_paragraph()
        run_graph_type = parag_image.add_run('Istogramma TOP tipologie trovate: ')
        run_graph_type.bold = True
        run_image = parag_image.add_run()
        run_image.add_picture(file_name,width=Inches(6.0), height=Inches(4.0))
        os.remove(file_name)

        parag_title_family = doc.add_paragraph()
        run_type = parag_title_family.add_run('TOP Famiglie di Malware trovati: ')
        run_type.bold = True
        
        list_number = 1
        sorted_malware,sorted_dict = find_top_malware(dict_malware_family,malware_family)
        file_name = malware_graph(sorted_dict,url_splitted[2],'family')

        for malware in sorted_malware:
            parag_family = doc.add_paragraph()
            run_list = parag_family.add_run('     '+str(list_number)+'. ')
            run_list.bold = True
            run_family = parag_family.add_run(malware+' trovato '+str(sorted_dict[malware])+' volte/a')
            list_number += 1

        parag_image = doc.add_paragraph()
        run_graph_fam = parag_image.add_run('Istogramma TOP famiglie trovate: ')
        run_graph_fam.bold = True
        run_image = parag_image.add_run()
        run_image.add_picture(file_name,width=Inches(6.0), height=Inches(4.5))
        os.remove(file_name)
        
        
    #se non sono stati trovati Malware lo scrive
    else:
        parag_malware = doc.add_paragraph()
        run_malware = parag_malware.add_run('NESSUN MALWARE TROVATO')
        run_malware.bold = True

    #salvataggio del documento
    doc.save(f'Reports/Report {url_splitted[2]}.docx')


'''
#questo metodo si occupa di scrivere tutte le informazioni utili per il report
#riguardo un'URL, all'interno di un file word
def write_report(url, whois_info, malware_info, malware_number,category_type, category_description):

    #creazione oggetti Whois e Malware
    whois_obj = Whois()
    malware_obj = Malware()
    
    #suddivisione dell'URL ricevuto, da utilizzare nel titolo 
    #e nel nome del file
    url_splitted = url.split('/')
    #impostazione generali del documento
    doc  = docx.Document()
    style = doc.styles['Normal']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    #scrittura del titolo
    heading = doc.add_paragraph()
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_heading = heading.add_run(f'Report del sito {url_splitted[2]}')
    run_heading.font.size = Pt(24)
    run_heading.font.color.rgb = RGBColor.from_string('000000')
    run_heading.bold = True

    #spaziatura 
    doc.add_paragraph()

    #acquisizione record WHOIS per scrittua d'informazioni
    whois_obj = whois_info

    #scrittua del dominio e categoria, dominio preso dal record WHOIS e categoria tramite X-Force
    parag_domain = doc.add_paragraph()
    run_domain = parag_domain.add_run('Dominio: ')
    run_domain.bold = True
    run_site_domain = parag_domain.add_run(str(whois_obj.get_domain()))
    parag_category = doc.add_paragraph()
    run_category = parag_category.add_run('Categoria: ')
    run_category.bold = True
    #se la categoria esiste la scrive, altrimenti no
    if category_type != True:
        parag_category.add_run(str(category_type)+': '+str(category_description))

    #spaziatura
    doc.add_paragraph()

    #scrittura di tutti i campi del record WHOIS
    parag_whois = doc.add_paragraph()
    run_whois = parag_whois.add_run('RECORD WHOIS:')
    run_whois.bold = True


    parag_creation = doc.add_paragraph()
    run_creation = parag_creation.add_run('Data di Creazione: ')
    run_creation.bold = True
    parag_creation.style = list_type
    parag_creation.add_run(str(whois_obj.get_creation_date()))

    parag_update = doc.add_paragraph()
    run_update = parag_update.add_run('Data di Aggiornamento: ')
    run_update.bold = True
    parag_update.style = list_type
    parag_update.add_run(str(whois_obj.get_updated_date()))

    parag_expiration = doc.add_paragraph()
    run_expiration = parag_expiration.add_run('Data di Scadenza: ')
    run_expiration.bold = True
    parag_expiration.style = list_type
    parag_expiration.add_run(str(whois_obj.get_expiration_date()))

    parag_name = doc.add_paragraph()
    run_name = parag_name.add_run('Nome del registrante: ')
    run_name.bold = True
    parag_name.style = list_type
    parag_name.add_run(str(whois_obj.get_name()))

    parag_org = doc.add_paragraph()
    run_org = parag_org.add_run('Organizzazione del registrante: ')
    run_org.bold = True
    parag_org.style = list_type
    parag_org.add_run(str(whois_obj.get_org()))

    parag_location = doc.add_paragraph()
    run_location = parag_location.add_run('Città e/o regione del registrante: ')
    run_location.bold = True
    parag_location.style = list_type
    parag_location.add_run(str(whois_obj.get_location()))

    parag_registrar = doc.add_paragraph()
    run_registrar = parag_registrar.add_run('Nome del registrar: ')
    run_registrar.bold = True
    parag_registrar.style = list_type
    parag_registrar.add_run(str(whois_obj.get_registrar()))

    parag_email = doc.add_paragraph()
    run_email = parag_email.add_run('Email associate: ')
    run_email.bold = True
    parag_email.style = list_type
    parag_email.add_run(str(whois_obj.get_email()))

    #spaziatura
    doc.add_paragraph()

    #se sono stati trovati malware per quell'URL
    if malware_info != True:
        #scrive quanti ne sono stati trovati
        parag_malware = doc.add_paragraph()
        run_malware = parag_malware.add_run(f'MALWARE TROVATI {malware_number}:')
        run_malware.bold = True
        number = 1
        #e scrive tutte le informazioni inerenti a ciascuno
        for malware in malware_info:
            
            malware_obj = malware

            parag_number = doc.add_paragraph()
            run_number = parag_number.add_run(f'     Numero {number}:')
            run_number.bold = True

            parag_domain = doc.add_paragraph()
            run_domain = parag_domain.add_run('Dominio: ')
            run_domain.bold = True
            parag_domain.style = list_type
            parag_domain.add_run(str(malware_obj.get_domain()))

            parag_type = doc.add_paragraph()
            run_type = parag_type.add_run('Tipologia Malware: ')
            run_type.bold = True
            parag_type.style = list_type
            parag_type.add_run(str(malware_obj.get_malware_type()))

            parag_family = doc.add_paragraph()
            run_family = parag_family.add_run('Famiglie del Malware: ')
            run_family.bold = True
            parag_family.style = list_type
            parag_family.add_run(str(malware_obj.get_family()))

            parag_file_path = doc.add_paragraph()
            run_file_path = parag_file_path.add_run('File Path: ')
            run_file_path.bold = True
            parag_file_path.style = list_type
            parag_file_path.add_run(str(malware_obj.get_filepath()))

            parag_ip = doc.add_paragraph()
            run_ip = parag_ip.add_run('IP: ')
            run_ip.bold = True
            parag_ip.style = list_type
            parag_ip.add_run(str(malware_obj.get_ip()))

            doc.add_paragraph()
            number += 1
    #se non sono stati trovati Malware lo scrive
    else:
        parag_malware = doc.add_paragraph()
        run_malware = parag_malware.add_run('NESSUN MALWARE TROVATO')
        run_malware.bold = True

    #salvataggio del documento
    doc.save(f'Reports/Report {url_splitted[2]}.docx')
'''


